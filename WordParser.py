import docx

d = docx.Document("demo.docx")
print(d.paragraphs)
print("\n")
print(d.paragraphs[0])
print("\n")
print(d.paragraphs[0].text)
print("\n")
print(d.paragraphs[1].text)
print("\n")

p = d.paragraphs[1]
print(p.runs)  # segment of text with a particular format
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)
print(p.runs[4].text)

print(p.runs[2].bold)
print(p.runs[4].italic)
p.runs[4].underline = True
p.runs[4].text = "italic and underlined."
d.save("demo2.docx")

# changing style
p.style = "Title"
d.save("demo3.docx")

# creating a new document
d = docx.Document()
d.add_paragraph("Hello this is a paragraph.")
d.add_paragraph("This is another paragraph.")
d.save("demo4.docx")
p = d.paragraphs[0]
p.add_run("This is a new run.")
print(p.runs)
p.runs[1].bold = True
d.save("demo5.docx")
